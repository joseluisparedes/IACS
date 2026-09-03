import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import os

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def create_test_document():
    doc = docx.Document()

    # Set page margins
    for s in doc.sections:
        s.top_margin = Inches(1)
        s.bottom_margin = Inches(1)
        s.left_margin = Inches(1)
        s.right_margin = Inches(1)

    NAVY = RGBColor(27, 54, 93)      # #1B365D
    INDIGO = RGBColor(79, 90, 245)   # #4F5AF5
    DARK = RGBColor(51, 65, 85)      # #334155
    SLATE = RGBColor(100, 116, 139)  # #64748B

    # Document Header
    p_pre = doc.add_paragraph()
    p_pre.paragraph_format.space_before = Pt(0)
    p_pre.paragraph_format.space_after = Pt(2)
    r_pre = p_pre.add_run("IACS · CASO DE PRUEBA DE BASE DE CONOCIMIENTO")
    r_pre.font.name = "Arial"
    r_pre.font.size = Pt(9.5)
    r_pre.font.bold = True
    r_pre.font.color.rgb = INDIGO

    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_after = Pt(6)
    r_title = p_title.add_run("Manual de Políticas y Estándares de TI para Iniciativas 2026")
    r_title.font.name = "Arial"
    r_title.font.size = Pt(18)
    r_title.font.bold = True
    r_title.font.color.rgb = NAVY

    p_desc = doc.add_paragraph()
    p_desc.paragraph_format.space_after = Pt(14)
    r_desc = p_desc.add_run("Documento con contenido ficticio institucional para probar la extracción, segmentación y enriquecimiento de contexto en el Asistente Teo.")
    r_desc.font.name = "Arial"
    r_desc.font.size = Pt(10.5)
    r_desc.font.italic = True
    r_desc.font.color.rgb = SLATE

    # Divider line
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_after = Pt(14)
    p_div_border = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="12" w:space="1" w:color="4F5AF5"/></w:pBdr>')
    p_div._p.get_or_add_pPr().append(p_div_border)

    def add_section_box(title, domain, content, rule, teo_action):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False
        cell = tbl.cell(0, 0)
        cell.width = Inches(6.5)
        set_cell_background(cell, "FFFFFF")
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)

        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="20" w:space="0" w:color="4F5AF5"/><w:top w:val="single" w:sz="6" w:space="0" w:color="E2E8F0"/><w:right w:val="single" w:sz="6" w:space="0" w:color="E2E8F0"/><w:bottom w:val="single" w:sz="6" w:space="0" w:color="E2E8F0"/></w:tcBorders>')
        tcPr.append(tcBorders)

        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r_t = p.add_run(f"📌 {title}\n")
        r_t.font.name = "Arial"
        r_t.font.size = Pt(12)
        r_t.font.bold = True
        r_t.font.color.rgb = NAVY

        p_dom = cell.add_paragraph()
        p_dom.paragraph_format.space_after = Pt(6)
        r_d = p_dom.add_run(f"Categoría: {domain}")
        r_d.font.name = "Arial"
        r_d.font.size = Pt(9)
        r_d.font.bold = True
        r_d.font.color.rgb = SLATE

        p_c = cell.add_paragraph()
        p_c.paragraph_format.space_after = Pt(6)
        r_ch = p_c.add_run("Contexto y Descripción: ")
        r_ch.font.bold = True
        r_ch.font.size = Pt(9.5)
        r_ch.font.color.rgb = DARK
        r_cb = p_c.add_run(content)
        r_cb.font.size = Pt(9.5)
        r_cb.font.color.rgb = DARK

        p_r = cell.add_paragraph()
        p_r.paragraph_format.space_after = Pt(6)
        r_rh = p_r.add_run("Regla Institucional Obligatoria: ")
        r_rh.font.bold = True
        r_rh.font.size = Pt(9.5)
        r_rh.font.color.rgb = NAVY
        r_rb = p_r.add_run(rule)
        r_rb.font.size = Pt(9.5)
        r_rb.font.color.rgb = DARK

        p_q = cell.add_paragraph()
        p_q.paragraph_format.space_after = Pt(2)
        r_qh = p_q.add_run("Cuestionamiento que debe aplicar TEO: ")
        r_qh.font.bold = True
        r_qh.font.size = Pt(9.5)
        r_qh.font.color.rgb = INDIGO
        r_qb = p_q.add_run(teo_action)
        r_qb.font.italic = True
        r_qb.font.size = Pt(9.5)
        r_qb.font.color.rgb = NAVY

        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 1. Pasarelas de Pago
    add_section_box(
        title="Política de Integración de Pasarelas de Pago y Cobranzas",
        domain="Finanzas y Sistemas Transaccionales",
        content=(
            "La institución cuenta con convenios corporativos homologados para recaudación digital con las pasarelas Niubiz y PayU. "
            "Todas las transacciones de matrícula, pensiones, certificados y trámites académicos deben procesarse exclusivamente a través de estas pasarelas. "
            "La conciliación bancaria se ejecuta en procesos batch programados cada 4 horas."
        ),
        rule=(
            "Está terminantemente prohibido contratar o solicitar pasarelas de pago independientes (como Stripe, PayPal o MercadoPago) "
            "sin la aprobación formal previa de la Dirección de Finanzas y la Gerencia de Ciberseguridad."
        ),
        teo_action=(
            "Si el solicitante pide implementar cobros o pagos en línea, TEO debe preguntar: "
            "'¿Has contemplado utilizar las pasarelas homologadas Niubiz o PayU? Recuerda que cualquier pasarela externa requiere aprobación de Finanzas y Ciberseguridad. "
            "¿Cuál es el volumen mensual aproximado de transacciones que esperas recaudar?'"
        )
    )

    # 2. Arquitectura de Integración Banner
    add_section_box(
        title="Estándar de Integración con el Sistema Académico Core (Banner)",
        domain="Arquitectura de TI e Integraciones",
        content=(
            "El sistema Banner es el repositorio único y maestro de la vida académica del estudiante (expedientes, notas, convalidaciones y mallas curriculares). "
            "Cualquier sistema satélite que requiera consultar o escribir información debe hacerlo mediante el API Gateway corporativo (WSO2) "
            "utilizando tokens de autenticación OAuth2 y servicios web REST autorizados."
        ),
        rule=(
            "Queda estrictamente prohibido solicitar conexiones JDBC/ODBC directas a la base de datos Oracle de producción de Banner. "
            "Tampoco se autorizan réplicas nocturnas completas de bases de datos para evitar desincronización de registros."
        ),
        teo_action=(
            "Si la iniciativa involucra alumnos, docentes o notas, TEO debe interrogar: "
            "'Dado que esta información reside en Banner: ¿la solución contempla consumir las APIs existentes del API Gateway institucional? "
            "Recuerda que no se permiten conexiones directas a la base de datos de Banner.'"
        )
    )

    # 3. Justificación Financiera y Retorno de Inversión
    add_section_box(
        title="Criterio de Viabilidad y Umbral Mínimo de Ahorro de Horas-Hombre",
        domain="Evaluación y Priorización de Demanda TI",
        content=(
            "Los recursos del equipo de TI están priorizados para proyectos de alto impacto institucional. "
            "Para que un proyecto de automatización sea aprobado por el comité, debe demostrar un beneficio mínimo de 40 horas-hombre mensuales "
            "o un beneficio económico directo superior a $3,000 USD anuales en reducción de pérdidas o multas regulatorias."
        ),
        rule=(
            "Requerimientos con un ahorro menor a 15 horas mensuales no califican como iniciativas de desarrollo en IACS. "
            "Dichas necesidades deben resolverse internamente por el área usuaria mediante herramientas de autoservicio como Power Automate o macros."
        ),
        teo_action=(
            "Si el usuario describe una tarea manual repetitiva, TEO debe indagar los números concretos: "
            "'Para calcular la prioridad de tu iniciativa ante el comité: ¿cuántos colaboradores intervienen en esta tarea y cuántas horas a la semana le dedican a este proceso manual? "
            "Necesitamos validar que supere el umbral mínimo de 40 horas mensuales.'"
        )
    )

    # 4. Seguridad de la Información y Single Sign-On (SSO)
    add_section_box(
        title="Política Obligatoria de Autenticación Centralizada (Azure AD / SSO)",
        domain="Ciberseguridad y Accesos",
        content=(
            "Todos los aplicativos web, portales de autoservicio y aplicaciones móviles institucionales deben integrarse "
            "al Single Sign-On (SSO) corporativo basado en Microsoft Entra ID (Azure Active Directory) con autenticación multifactor (MFA)."
        ),
        rule=(
            "No se aceptará ninguna iniciativa técnica que proponga la creación de una base de datos propia de usuarios, "
            "contraseñas locales o mecanismos de registro independientes que no utilicen el correo institucional (@upn.pe o @laureate.net)."
        ),
        teo_action=(
            "En la etapa de arquitectura y pruebas, TEO debe validar: "
            "'¿El acceso de los usuarios estará 100% integrado al SSO institucional con Microsoft Entra ID (Azure AD)? "
            "¿Se requiere definir roles específicos como solo lectura, aprobador o administrador?'"
        )
    )

    out_path = os.path.abspath("Documento_Prueba_Base_Conocimiento_IACS.docx")
    doc.save(out_path)
    print(f"Document created successfully at: {out_path}")

if __name__ == "__main__":
    create_test_document()
