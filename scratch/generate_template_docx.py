import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import os

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def create_document():
    doc = docx.Document()

    # Configure Margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    NAVY = RGBColor(27, 54, 93)       # #1B365D
    INDIGO = RGBColor(79, 90, 245)    # #4F5AF5
    DARK_GRAY = RGBColor(51, 65, 85)  # #334155
    SLATE = RGBColor(100, 116, 139)   # #64748B

    # Document Title / Header Box
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run_pre = title_p.add_run("IACS · PLATAFORMA DE GESTIÓN DE INICIATIVAS DE TI\n")
    run_pre.font.name = "Arial"
    run_pre.font.size = Pt(10)
    run_pre.font.bold = True
    run_pre.font.color.rgb = INDIGO

    run_title = title_p.add_run("Plantilla Oficial y Guía de Carga: Base de Conocimiento para TEO")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = NAVY

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(18)
    run_sub = sub_p.add_run("Guía colaborativa para estructurar fichas de contexto que potencien la capacidad de análisis y cuestionamiento del Asistente de IA.")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(11)
    run_sub.font.italic = True
    run_sub.font.color.rgb = SLATE

    # Horizontal divider
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_after = Pt(14)
    p_div_border = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="12" w:space="1" w:color="4F5AF5"/></w:pBdr>')
    p_div._p.get_or_add_pPr().append(p_div_border)

    # 1. OBJETIVO Y REGLA DE ORO
    h1 = doc.add_heading(level=1)
    h1_run = h1.add_run("1. Propósito y Regla de Oro para Cargar Conocimiento")
    h1_run.font.name = "Arial"
    h1_run.font.color.rgb = NAVY
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(6)

    p_intro = doc.add_paragraph(
        "El asistente TEO actúa como un Analista de Negocio Senior de TI (Senior Business Analyst). "
        "Su objetivo no es ser un formulario pasivo, sino cuestionar, validar y desafiar constructivamente al solicitante "
        "para que la iniciativa llegue al comité de evaluación con rigor técnico, justificación cuantitativa y alineación estratégica."
    )
    p_intro.paragraph_format.space_after = Pt(8)

    # Callout Box: Regla de Oro
    callout = doc.add_table(rows=1, cols=1)
    callout.alignment = WD_TABLE_ALIGNMENT.CENTER
    callout.autofit = False
    c_cell = callout.cell(0, 0)
    c_cell.width = Inches(6.5)
    set_cell_background(c_cell, "F1F5F9")  # Slate 100
    set_cell_margins(c_cell, top=140, bottom=140, left=180, right=180)

    # Left border highlight
    tcPr = c_cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="0" w:color="4F5AF5"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
    tcPr.append(tcBorders)

    cp = c_cell.paragraphs[0]
    cp.paragraph_format.space_after = Pt(4)
    c_bold = cp.add_run("📌 LA REGLA DE ORO: FICHAS ATÓMICAS (NO MANUALES DE 50 PÁGINAS)\n")
    c_bold.font.name = "Arial"
    c_bold.font.size = Pt(10.5)
    c_bold.font.bold = True
    c_bold.font.color.rgb = NAVY

    c_text = cp.add_run(
        "Para no saturar la memoria ni los tokens de la IA en cada conversación, NO subas manuales extensos sin procesar. "
        "El conocimiento debe cargarse en 'Fichas Atómicas de Decisión' (150 a 300 palabras). Cada ficha debe responder: "
        "¿Qué existe o cuál es la norma?, ¿Cuándo aplica? y ¿Qué pregunta o validación debe exigir TEO?"
    )
    c_text.font.name = "Arial"
    c_text.font.size = Pt(9.5)
    c_text.font.color.rgb = DARK_GRAY

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 2. ESTRUCTURA DE UNA FICHA
    h2 = doc.add_heading(level=1)
    h2_run = h2.add_run("2. Estructura Estándar de una Ficha de Conocimiento")
    h2_run.font.name = "Arial"
    h2_run.font.color.rgb = NAVY
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(6)

    table_struct = doc.add_table(rows=6, cols=2)
    table_struct.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_struct.autofit = False

    headers = [("Campo", Inches(1.8)), ("Descripción y Qué debe contener", Inches(4.7))]
    rows_data = [
        ("Título de la Ficha", "Nombre claro y conciso del sistema, política o flujo (ej. 'Sistema Core: Banner Académico')."),
        ("Dominio / Categoría", "Área temática: Sistemas Core | Evaluación Financiera | Políticas TI | Admisiones | etc."),
        ("Palabras Clave (Triggers)", "Términos que hacen que TEO recuerde esta ficha (ej. 'matrícula, notas, pagos, ahorro, horas')."),
        ("Descripción y Regla", "¿Qué existe actualmente y qué está estrictamente permitido o prohibido?"),
        ("Acción y Cuestionamiento de TEO", "La pregunta o desafío exacto que TEO debe hacer al solicitante al tocar este tema.")
    ]

    # Style Header
    hdr_cells = table_struct.rows[0].cells
    for i, (title, width) in enumerate(headers):
        hdr_cells[i].width = width
        set_cell_background(hdr_cells[i], "1B365D")
        set_cell_margins(hdr_cells[i], top=100, bottom=100, left=120, right=120)
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(title)
        run.font.name = "Arial"
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)

    for row_idx, (field, desc) in enumerate(rows_data, start=1):
        row = table_struct.rows[row_idx]
        c0, c1 = row.cells[0], row.cells[1]
        c0.width = Inches(1.8)
        c1.width = Inches(4.7)
        bg = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
        set_cell_background(c0, bg)
        set_cell_background(c1, bg)
        set_cell_margins(c0, top=80, bottom=80, left=120, right=120)
        set_cell_margins(c1, top=80, bottom=80, left=120, right=120)
        
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(field)
        r0.font.name = "Arial"
        r0.font.bold = True
        r0.font.size = Pt(9.5)
        r0.font.color.rgb = NAVY

        p1 = c1.paragraphs[0]
        r1 = p1.add_run(desc)
        r1.font.name = "Arial"
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = DARK_GRAY

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 3. EJEMPLOS REALES RESUELTOS
    h3 = doc.add_heading(level=1)
    h3_run = h3.add_run("3. Ejemplos Reales Resueltos (Modelos de Referencia)")
    h3_run.font.name = "Arial"
    h3_run.font.color.rgb = NAVY
    h3.paragraph_format.space_before = Pt(14)
    h3.paragraph_format.space_after = Pt(6)

    def add_example_card(title, domain, triggers, desc, rules, questions):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False
        cell = tbl.cell(0, 0)
        cell.width = Inches(6.5)
        set_cell_background(cell, "FFFFFF")
        set_cell_margins(cell, top=120, bottom=120, left=140, right=140)

        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="18" w:space="0" w:color="4F5AF5"/><w:top w:val="single" w:sz="6" w:space="0" w:color="E2E8F0"/><w:right w:val="single" w:sz="6" w:space="0" w:color="E2E8F0"/><w:bottom w:val="single" w:sz="6" w:space="0" w:color="E2E8F0"/></w:tcBorders>')
        tcPr.append(tcBorders)

        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r_t = p.add_run(f"📋 {title}\n")
        r_t.font.name = "Arial"
        r_t.font.bold = True
        r_t.font.size = Pt(11)
        r_t.font.color.rgb = NAVY

        p_meta = cell.add_paragraph()
        p_meta.paragraph_format.space_after = Pt(4)
        r_m1 = p_meta.add_run(f"Dominio: {domain} | ")
        r_m1.font.bold = True
        r_m1.font.size = Pt(9)
        r_m1.font.color.rgb = SLATE
        r_m2 = p_meta.add_run(f"Triggers: {triggers}")
        r_m2.font.italic = True
        r_m2.font.size = Pt(9)
        r_m2.font.color.rgb = SLATE

        p_desc = cell.add_paragraph()
        p_desc.paragraph_format.space_after = Pt(4)
        r_dh = p_desc.add_run("Descripción: ")
        r_dh.font.bold = True
        r_dh.font.size = Pt(9.5)
        r_dh.font.color.rgb = DARK_GRAY
        r_db = p_desc.add_run(desc)
        r_db.font.size = Pt(9.5)
        r_db.font.color.rgb = DARK_GRAY

        p_rules = cell.add_paragraph()
        p_rules.paragraph_format.space_after = Pt(4)
        r_rh = p_rules.add_run("Regla Institucional: ")
        r_rh.font.bold = True
        r_rh.font.size = Pt(9.5)
        r_rh.font.color.rgb = DARK_GRAY
        r_rb = p_rules.add_run(rules)
        r_rb.font.size = Pt(9.5)
        r_rb.font.color.rgb = DARK_GRAY

        p_q = cell.add_paragraph()
        p_q.paragraph_format.space_after = Pt(2)
        r_qh = p_q.add_run("Cuestionamiento de TEO: ")
        r_qh.font.bold = True
        r_qh.font.size = Pt(9.5)
        r_qh.font.color.rgb = INDIGO
        r_qb = p_q.add_run(questions)
        r_qb.font.italic = True
        r_qb.font.size = Pt(9.5)
        r_qb.font.color.rgb = NAVY

        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Example 1
    add_example_card(
        "Sistema Core: ERP Académico (Banner)",
        "Sistemas Core",
        "matrícula, notas, profesores, horarios, alumnos, historial académico",
        "Banner es el ERP académico institucional donde residen los expedientes de estudiantes, planes de estudio y calificaciones.",
        "Está prohibido autorizar desarrollos satélite independientes para registro de notas o matrículas; cualquier necesidad debe solicitarse como módulo o integración sobre Banner.",
        "Si el usuario pide un portal de notas o inscripciones, TEO debe cuestionar: 'La institución ya cuenta con Banner para esta función. ¿Tu iniciativa busca integrarse con Banner o por qué motivo no se puede resolver con la funcionalidad nativa actual?'"
    )

    # Example 2
    add_example_card(
        "Filtro de Demanda: Criterio de Viabilidad de Proyecto TI",
        "Gobernanza TI",
        "laptops, monitores, licencias office, soporte, acceso, contraseña, ticket",
        "IACS gestiona iniciativas de desarrollo, mejora de procesos e integración tecnológica. No gestiona compras de consumibles ni tickets de soporte.",
        "Reemplazos de hardware, accesos a carpetas compartidas y fallos de sistemas existentes deben canalizarse por la Mesa de Ayuda (Service Desk).",
        "Si el usuario pide equipamiento o soporte de incidentes, TEO debe responder: 'Esta solicitud corresponde a soporte operativo. Debes gestionarla por Service Desk. IACS está destinado a proyectos e iniciativas de mejora tecnológica.'"
    )

    # Example 3
    add_example_card(
        "Métricas: Justificación Cuantitativa de Ahorro",
        "Evaluación Financiera",
        "ahorro, tiempo, horas hombre, manual, automatizar, personas, dinero",
        "Para sustentar la prioridad de una iniciativa, se requiere calcular el ahorro mensual: (N° de Personas) × (Horas semanales dedicadas) × 4 semanas.",
        "No se admiten respuestas vagas como 'ahorraremos mucho tiempo'. La iniciativa debe registrar horas-hombre concretas o beneficio económico.",
        "Si el usuario indica que ahorrará tiempo, TEO debe cuestionar: 'Para defender la iniciativa ante el comité: ¿cuántos colaboradores realizan hoy esta tarea manual y cuántas horas a la semana le dedica cada uno? Calculemos el ahorro de horas mensual.'"
    )

    # 4. PLANTILLAS EN BLANCO PARA COMPLETAR
    doc.add_page_break()
    h4 = doc.add_heading(level=1)
    h4_run = h4.add_run("4. Plantillas Vacías para Completar por tu Equipo")
    h4_run.font.name = "Arial"
    h4_run.font.color.rgb = NAVY
    h4.paragraph_format.space_before = Pt(14)
    h4.paragraph_format.space_after = Pt(6)

    p_fill = doc.add_paragraph("Completa los siguientes bloques con las políticas, sistemas o procesos de tu área para subirlos a la Base de Conocimiento de TEO:")
    p_fill.paragraph_format.space_after = Pt(10)

    for i in range(1, 4):
        p_t = doc.add_paragraph()
        p_t.paragraph_format.space_before = Pt(8)
        p_t.paragraph_format.space_after = Pt(2)
        r = p_t.add_run(f"Ficha N° {i} [Escribe el Título]")
        r.font.name = "Arial"
        r.font.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = INDIGO

        tbl_empty = doc.add_table(rows=5, cols=2)
        tbl_empty.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl_empty.autofit = False

        f_fields = [
            ("Dominio / Área", "[Ej. Sistemas Core, Finanzas, Admisiones, TI]"),
            ("Palabras Clave (Triggers)", "[Términos que activan esta ficha separados por coma]"),
            ("Descripción / Qué existe", "[Describe el sistema, proceso o contexto actual]"),
            ("Regla Institucional", "[Qué está permitido, qué está prohibido o límites técnicos]"),
            ("Pregunta / Cuestionamiento de TEO", "[Qué debe preguntar o validar TEO ante el solicitante]")
        ]

        for row_i, (f_name, f_ph) in enumerate(f_fields):
            r = tbl_empty.rows[row_i]
            c0, c1 = r.cells[0], r.cells[1]
            c0.width = Inches(2.0)
            c1.width = Inches(4.5)
            set_cell_background(c0, "F1F5F9")
            set_cell_background(c1, "FFFFFF")
            set_cell_margins(c0, top=80, bottom=80, left=100, right=100)
            set_cell_margins(c1, top=80, bottom=80, left=100, right=100)

            p0 = c0.paragraphs[0]
            rp0 = p0.add_run(f_name)
            rp0.font.name = "Arial"
            rp0.font.bold = True
            rp0.font.size = Pt(9.5)
            rp0.font.color.rgb = NAVY

            p1 = c1.paragraphs[0]
            rp1 = p1.add_run(f_ph)
            rp1.font.name = "Arial"
            rp1.font.italic = True
            rp1.font.size = Pt(9.5)
            rp1.font.color.rgb = SLATE

        doc.add_paragraph().paragraph_format.space_after = Pt(10)

    out_path = os.path.abspath("Plantilla_Base_Conocimiento_IACS_TEO.docx")
    doc.save(out_path)
    print(f"Document successfully created at: {out_path}")

if __name__ == "__main__":
    create_document()
